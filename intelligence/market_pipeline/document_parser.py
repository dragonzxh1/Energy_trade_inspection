"""Deterministic document parsing for the structured market pipeline."""

from __future__ import annotations

import html
import re
import uuid
from dataclasses import dataclass
from datetime import date, datetime
from html.parser import HTMLParser
from pathlib import Path
from typing import Iterable

from .contracts import (
    SOURCE_DOCUMENT_SCHEMA_VERSION,
    DateCandidate,
    DateCandidateSource,
    DocumentProcessingStatus,
    DocumentSection,
    ParseMethod,
    ParsedTable,
    SourceContent,
    SourceDocument,
    SourceDocumentMetadata,
    SourceIngestion,
    SourceStatus,
    TelegramInput,
)


MIN_PDF_TEXT_CHARS = 200
SOURCE_NAMESPACE = uuid.UUID("23aca9e4-8123-4d09-b55b-d45f35660792")
PUBLISHERS = (
    ("Platts", ("s&p global commodity insights", "platts"), 0.99),
    ("Argus", ("argus media", "argus"), 0.98),
    ("Reuters", ("reuters",), 0.98),
    ("Financial Times", ("financial times", "ft"), 0.98),
    ("The New York Times", ("new york times", "nyt"), 0.98),
    ("The Guardian", ("the guardian",), 0.98),
    ("The Wall Street Journal", ("wall street journal", "wsj"), 0.98),
    ("The Washington Post", ("washington post",), 0.98),
    ("Haaretz", ("haaretz",), 0.98),
    ("International Energy Agency", ("international energy agency", "iea"), 0.96),
    ("International Monetary Fund", ("international monetary fund", "imf"), 0.96),
    ("International Gas Union", ("international gas union", "igu"), 0.96),
)
REGIONS = ("Asia", "Europe", "Mediterranean", "Middle East", "Americas", "Atlantic", "Pacific")
COMMODITIES = (
    "crude oil", "crude", "naphtha", "gasoline", "gasoil", "diesel", "jet fuel",
    "fuel oil", "lng", "natural gas", "lpg", "ethylene", "propylene", "benzene",
)


@dataclass(frozen=True)
class PageText:
    page_number: int | None
    text: str


class _HTMLTextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.parts: list[str] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"p", "div", "br", "li", "h1", "h2", "h3", "h4", "tr"}:
            self.parts.append("\n")

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def text(self) -> str:
        return _clean_text(html.unescape("".join(self.parts)))


def _clean_text(value: str) -> str:
    value = value.replace("\x00", "")
    value = re.sub(r"[ \t]+", " ", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def _stable_id(prefix: str, seed: str) -> str:
    return f"{prefix}-{uuid.uuid5(SOURCE_NAMESPACE, seed)}"


def _convert_tables(
    tables: Iterable[tuple[int | None, list[list[str | None]]]], parse_method: str
) -> list[ParsedTable]:
    converted: list[ParsedTable] = []
    for index, (page_number, matrix) in enumerate(tables):
        if not matrix:
            continue
        width = max(len(row) for row in matrix)
        header = [str(value or "").strip() or f"column_{column + 1}" for column, value in enumerate(matrix[0])]
        header += [f"column_{column + 1}" for column in range(len(header), width)]
        rows = []
        for matrix_row in matrix[1:]:
            values = list(matrix_row) + [None] * (width - len(matrix_row))
            rows.append({header[column]: (str(values[column]).strip() if values[column] is not None else None) for column in range(width)})
        converted.append(
            ParsedTable(
                table_id=f"pending-{index}", source_id="pending", table_index=index,
                page_number=page_number, columns=header, rows=rows,
                parse_method=parse_method, parse_confidence=0.9,
            )
        )
    return converted


def _read_pdf(path: Path, *, extract_tables: bool) -> tuple[list[PageText], list[ParsedTable]]:
    import fitz

    pages: list[PageText] = []
    raw_tables: list[tuple[int, list[list[str | None]]]] = []
    with fitz.open(str(path)) as document:
        for page_number, page in enumerate(document, start=1):
            page_text = _clean_text(page.get_text("text"))
            pages.append(PageText(page_number, page_text))
            if not extract_tables or not _looks_like_price_table_page(page_text):
                continue
            try:
                finder = page.find_tables()
                for table in finder.tables:
                    raw_tables.append((page_number, table.extract()))
            except (AttributeError, RuntimeError, ValueError):
                continue
    return pages, _convert_tables(raw_tables, "pdf_table")


def _looks_like_price_table_page(text: str) -> bool:
    lowered = text.lower()
    labels = sum(
        marker in lowered
        for marker in ("assessment", "bid", "offer", "low", "high", "change", "midpoint", "$/")
    )
    numeric_rows = sum(
        1 for line in text.splitlines()
        if len(re.findall(r"(?<!\w)[+-]?\d+(?:\.\d+)?", line)) >= 3
    )
    return labels >= 2 and numeric_rows >= 2


def _supports_structured_pdf_tables(filename: str) -> bool:
    lowered = filename.lower()
    return any(
        marker in lowered
        for marker in ("marketscan", "oilgram", "bunkerwire", "tankerwire", "lpgaswire", "price report")
    )


def _read_docx(path: Path) -> tuple[list[PageText], list[ParsedTable]]:
    from docx import Document

    document = Document(str(path))
    text = _clean_text("\n".join(paragraph.text for paragraph in document.paragraphs))
    tables = [(None, [[cell.text for cell in row.cells] for row in table.rows]) for table in document.tables]
    return [PageText(None, text)], _convert_tables(tables, "docx_table")


def _parse_date(raw: str) -> date | None:
    normalized = raw.strip().replace("/", "-").replace(".", "-")
    for pattern in ("%Y-%m-%d", "%d-%m-%Y", "%m-%d-%Y", "%d %B %Y", "%B %d %Y", "%d %b %Y", "%b %d %Y"):
        try:
            return datetime.strptime(normalized.replace(",", ""), pattern).date()
        except ValueError:
            continue
    compact = re.fullmatch(r"(20\d{2})(\d{2})(\d{2})", normalized)
    if compact:
        try:
            return date(int(compact[1]), int(compact[2]), int(compact[3]))
        except ValueError:
            return None
    return None


DATE_PATTERN = re.compile(
    r"\b(?:20\d{2}[-/.]\d{1,2}[-/.]\d{1,2}|\d{1,2}[-/.]\d{1,2}[-/.]20\d{2}|"
    r"\d{1,2}\s+(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|"
    r"Jul(?:y)?|Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+20\d{2}|"
    r"(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:tember)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+\d{1,2},?\s+20\d{2})\b",
    re.IGNORECASE,
)


def _candidate(source: DateCandidateSource, evidence: str, confidence: float) -> DateCandidate | None:
    match = DATE_PATTERN.search(evidence)
    raw = match.group(0) if match else ""
    if not raw:
        compact = re.search(r"(?:20\d{6}|\d{8})", evidence)
        raw = compact.group(0) if compact else ""
    if re.fullmatch(r"\d{8}", raw) and not raw.startswith("20"):
        raw = f"{raw[4:8]}-{raw[2:4]}-{raw[0:2]}"
    parsed = _parse_date(raw) if raw else None
    if not parsed:
        return None
    return DateCandidate(value=parsed, source=source, evidence=evidence[:240], confidence=confidence)


def collect_date_candidates(text: str, filename: str, telegram_date: datetime) -> list[DateCandidate]:
    candidates: list[DateCandidate] = []
    assessment = re.search(r"(?im)^.{0,40}(?:assessment|assessed)\s+date.{0,80}$", text)
    if assessment:
        found = _candidate(DateCandidateSource.ASSESSMENT, assessment.group(0), 0.99)
        if found:
            candidates.append(found)
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    title_candidate = _candidate(DateCandidateSource.TITLE, " | ".join(lines[:5]), 0.94)
    if title_candidate:
        candidates.append(title_candidate)
    issue_header = re.search(
        r"(?im)^.{0,80}(?:volume\s+\d+\s*/\s*issue\s+\d+).{0,100}$",
        text,
    )
    body_candidate = (
        _candidate(DateCandidateSource.BODY, issue_header.group(0), 0.92)
        if issue_header else _candidate(DateCandidateSource.BODY, "\n".join(lines[5:80]), 0.82)
    )
    if body_candidate:
        candidates.append(body_candidate)
    published = re.search(r"(?im)^.{0,30}published(?:\s+at|\s+on)?[: ]+.{0,80}$", text)
    if published:
        published_candidate = _candidate(DateCandidateSource.PUBLISHED_AT, published.group(0), 0.76)
        if published_candidate:
            candidates.append(published_candidate)
    filename_candidate = _candidate(DateCandidateSource.FILENAME, filename, 0.7)
    if not filename_candidate:
        short_date = re.search(r"(?<!\d)(\d{2})(\d{2})(\d{2})?(?!\d)", filename)
        if short_date:
            day, month = int(short_date[1]), int(short_date[2])
            year = 2000 + int(short_date[3]) if short_date[3] else telegram_date.year
            try:
                filename_candidate = DateCandidate(
                    value=date(year, month, day), source=DateCandidateSource.FILENAME,
                    evidence=filename[:240], confidence=0.66 if short_date[3] else 0.58,
                )
            except ValueError:
                filename_candidate = None
    if filename_candidate:
        candidates.append(filename_candidate)
    candidates.append(
        DateCandidate(
            value=telegram_date.date(), source=DateCandidateSource.TELEGRAM,
            evidence=telegram_date.isoformat(), confidence=0.55,
        )
    )
    priority = {source: index for index, source in enumerate(DateCandidateSource)}
    return sorted(candidates, key=lambda item: priority[item.source])


def _detect_publisher(filename: str, text: str, forwarded_from: str | None) -> tuple[str, float]:
    normalized_filename = re.sub(
        r"(?<=[a-z])(?=\d)|(?<=\d)(?=[a-z])", " ", filename.lower()
    )
    haystacks = ((normalized_filename, -0.01), (text[:5000].lower(), -0.08), ((forwarded_from or "").lower(), -0.25))
    matches: list[tuple[float, str]] = []
    for publisher, aliases, confidence in PUBLISHERS:
        for haystack, penalty in haystacks:
            if any(
                bool(re.search(rf"\b{re.escape(alias)}\b", haystack)) if len(alias) <= 3 else alias in haystack
                for alias in aliases
            ):
                matches.append((max(0.5, confidence + penalty), publisher))
    if matches:
        confidence, publisher = max(matches)
        return publisher, confidence
    return "Unknown", 0.2


def _select_market_date(
    candidates: list[DateCandidate], telegram_date: datetime
) -> tuple[DateCandidate, str]:
    assessment = next((item for item in candidates if item.source == DateCandidateSource.ASSESSMENT), None)
    if assessment:
        return assessment, f"selected {assessment.source.value}: {assessment.evidence}"
    filename = next((item for item in candidates if item.source == DateCandidateSource.FILENAME), None)
    if filename and abs((filename.value - telegram_date.date()).days) <= 45:
        conflicting = [
            item for item in candidates
            if item.source in {DateCandidateSource.TITLE, DateCandidateSource.BODY, DateCandidateSource.PUBLISHED_AT}
            and abs((item.value - filename.value).days) > 45
        ]
        eligible = [item for item in candidates if item not in conflicting]
        selected = eligible[0]
        reason = f"selected {selected.source.value}: {selected.evidence}"
        if conflicting:
            reason += "; rejected stale/conflicting document date against filename issue date"
        return selected, reason
    selected = candidates[0]
    return selected, f"selected {selected.source.value}: {selected.evidence}"


def _detect_report_family(filename: str, publisher: str) -> str:
    lowered = filename.lower()
    families = (
        ("Asia-Pacific Arab Gulf Marketscan", ("asia-pacific", "arab gulf marketscan")),
        ("European Marketscan", ("european marketscan",)),
        ("US Marketscan", ("us marketscan",)),
        ("Oilgram Price Report", ("oilgram price report",)),
        ("Bunkerwire", ("bunkerwire",)),
        ("Clean Tankerwire", ("clean tankerwire",)),
        ("Dirty Tankerwire", ("dirty tankerwire",)),
        ("LPGaswire", ("lpgaswire",)),
        ("LNG Daily", ("lng daily",)),
        ("Gas Market Report", ("gas market report",)),
    )
    for family, markers in families:
        if all(marker in lowered for marker in markers):
            return family
    return publisher if publisher != "Unknown" else "unknown"


def _classify(text: str) -> tuple[list[str], list[str]]:
    lowered = text.lower()
    regions = [region for region in REGIONS if region.lower() in lowered]
    commodities = [commodity for commodity in COMMODITIES if re.search(rf"\b{re.escape(commodity)}\b", lowered)]
    return regions, commodities


def _looks_like_heading(line: str) -> bool:
    stripped = line.strip(" :-")
    if not stripped or len(stripped) > 120 or len(stripped.split()) > 14:
        return False
    lowered = stripped.casefold()
    if (
        lowered.startswith(("page ", "www.", "http://", "https://"))
        or "@" in stripped
        or stripped.endswith((".", ";", ","))
    ):
        return False
    numbered = re.match(r"^(\d+)(?:\.\d+){0,3}\.?\s+[A-Z]", stripped)
    if numbered and not re.fullmatch(r"20\d{2}", numbered.group(1)):
        return True
    known = (
        "abstract", "executive summary", "market review", "market outlook",
        "key developments", "highlights", "supply outlook", "demand outlook",
    )
    if lowered in known:
        return True
    classified = any(term.lower() in lowered for term in (*REGIONS, *COMMODITIES))
    styled = stripped.isupper() or stripped.istitle()
    if classified and styled:
        return True
    words = stripped.split()
    capitalized = sum(word[:1].isupper() for word in words if word)
    return len(words) >= 2 and capitalized / len(words) >= 0.9 and not re.search(r"[!?]", stripped)


def _split_sections(source_id: str, pages: list[PageText]) -> list[DocumentSection]:
    chunks: list[tuple[str, int | None, list[str]]] = []
    for page in pages:
        current_title = f"Page {page.page_number}" if page.page_number else "Document"
        current_lines: list[str] = []
        for line in (item.strip() for item in page.text.splitlines()):
            if re.fullmatch(r"(?:page\s*[|:]?\s*)?\d+", line, flags=re.IGNORECASE) or line.casefold().startswith("www."):
                continue
            if _looks_like_heading(line) and current_lines:
                chunks.append((current_title, page.page_number, current_lines))
                current_title, current_lines = line, []
            elif _looks_like_heading(line):
                current_title = line
            elif line:
                current_lines.append(line)
        if current_lines:
            chunks.append((current_title, page.page_number, current_lines))
    if not chunks:
        chunks = [("Document", None, [])]
    sections: list[DocumentSection] = []
    for index, (title, page_number, lines) in enumerate(chunks):
        section_text = _clean_text("\n".join(lines))
        regions, commodities = _classify(f"{title}\n{section_text}")
        sections.append(
            DocumentSection(
                section_id=_stable_id("SEC", f"{source_id}:{index}:{title}"), source_id=source_id,
                section_index=index, section_title=title, page_start=page_number, page_end=page_number,
                region=regions[0] if regions else None, commodity=commodities[0] if commodities else None,
                section_type="market_commentary" if section_text else "page", text=section_text,
                classification_confidence=0.9 if regions or commodities else 0.55,
            )
        )
    return sections


def parse_telegram_document(
    telegram_input: TelegramInput,
    *,
    attachment_id: str | None = None,
    parsed_text_dir: Path | None = None,
    parser_version: str = SOURCE_DOCUMENT_SCHEMA_VERSION,
) -> SourceDocument:
    attachment = telegram_input.attachment
    path = Path(attachment.attachment_path)
    source_id = _stable_id("SRC", f"{attachment.attachment_hash.lower()}:{parser_version}")
    review_reasons: list[str] = []
    pages: list[PageText] = []
    tables: list[ParsedTable] = []
    method = ParseMethod.UNSUPPORTED
    error_message: str | None = None
    try:
        suffix = path.suffix.lower()
        mime = attachment.attachment_mime_type.lower()
        if mime.startswith("image/") or suffix in {".png", ".jpg", ".jpeg", ".webp", ".tiff"}:
            method = ParseMethod.IMAGE_ONLY
            review_reasons.append("image-only attachment; production OCR is disabled")
        elif suffix == ".pdf" or mime == "application/pdf":
            pages, tables = _read_pdf(path, extract_tables=_supports_structured_pdf_tables(attachment.attachment_name))
            method = ParseMethod.PDF_TEXT
            if len("\n".join(page.text for page in pages).strip()) < MIN_PDF_TEXT_CHARS:
                review_reasons.append(f"image-only or low-text PDF (< {MIN_PDF_TEXT_CHARS} chars)")
        elif suffix == ".docx" or "wordprocessingml" in mime:
            pages, tables = _read_docx(path)
            method = ParseMethod.PLAIN_TEXT
        elif suffix in {".html", ".htm"} or mime == "text/html":
            extractor = _HTMLTextExtractor()
            extractor.feed(path.read_text(encoding="utf-8", errors="replace"))
            pages = [PageText(None, extractor.text())]
            method = ParseMethod.HTML
        elif suffix in {".txt", ".md"} or mime.startswith("text/"):
            pages = [PageText(None, _clean_text(path.read_text(encoding="utf-8", errors="replace")))]
            method = ParseMethod.PLAIN_TEXT
        else:
            review_reasons.append(f"unsupported document format: {suffix or mime}")
    except Exception as error:
        error_message = str(error)
        review_reasons.append(f"parse failed: {error}")

    parsed_text = _clean_text("\n\n".join(page.text for page in pages if page.text))
    publisher, publisher_confidence = _detect_publisher(
        attachment.attachment_name, parsed_text, telegram_input.message.forwarded_from
    )
    report_family = _detect_report_family(attachment.attachment_name, publisher)
    if publisher == "Unknown":
        review_reasons.append("publisher could not be verified")
    date_candidates = collect_date_candidates(parsed_text, attachment.attachment_name, telegram_input.message.telegram_message_date)
    selected_date, market_date_reason = _select_market_date(
        date_candidates, telegram_input.message.telegram_message_date
    )
    published_candidate = next(
        (candidate for candidate in date_candidates if candidate.source == DateCandidateSource.PUBLISHED_AT),
        None,
    )
    published_at = (
        datetime.combine(published_candidate.value, datetime.min.time(), tzinfo=telegram_input.message.telegram_message_date.tzinfo)
        if published_candidate else None
    )
    lines = [line.strip() for line in parsed_text.splitlines() if line.strip()]

    # Patterns that indicate a line is NOT a meaningful title
    SKIP_TITLE_PATTERN = re.compile(
        r"(?ix)"
        r"^\s*\(?\s*(?:continued\s+(?:on|from)\s+page\s*\d+\s*\)?\s*)$|"
        r"^\s*(?:https?://|www\.)\S+|"
        r"^\s*Volume\s+\d+\s*/\s*Issue|"
        r"^\s*Page\s+\d+|"
        r"^\s*\d+\s*$|"
        r"^\s*[*•●‣◦⁃∙\s]+\s*$|"  # lines of only symbols/bullets
        r"^\s*[A-Za-z]+\d{2,6}[A-Za-z]*\s*$|"  # bare filename: FT2107US
        r"^\s*(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}\s*$|"  # bare month year
        r"^\s*photo[_\s]+\d+|\w+_\d{6,8}_\d{6}\s*$|"  # photo filenames
        r"^\s*$"
    )
    # Lines that look like publication/section names
    PUBLICATION_NAME_PATTERN = re.compile(
        r"(?i)^(?:Platts|S&P\s*Global|Argus|Reuters|Bloomberg|ICE|CME|OPEC|EIA|IEA)\b"
    )
    # Tier 1b: Known report/publication names without publisher prefix
    KNOWN_REPORT_PATTERN = re.compile(
        r"(?i)^(?:Oilgram\s+Price\s+Report|"
        r"Asia-Pacific\s+Arab\s+Gulf\s+Marketscan|European\s+Marketscan|US\s+Marketscan|"
        r"Bunkerwire|Clean\s+Tankerwire|Dirty\s+Tankerwire|LPGaswire|"
        r"LNG\s+Daily|Gas\s+Market\s+Report|"
        r"(?:Weekly|Monthly|Daily|Annual)\s+[\w\s]+(?:Report|Update|Review|Outlook)|\b"
        r"Global\s+Critical\s+Minerals|"
        r"(?:The\s+)?Economist|Barron's|Forbes|Bloomberg\s+Businessweek)\b"
    )
    # Tier 1.5: Title-case lines that look like proper titles (3-12 words)
    TITLE_LIKE_PATTERN = re.compile(
        r"^(?:[A-Z][a-z]+(?:[\s\-][A-Z][a-z]+){2,11})$"
    )

    # Tier 1: Find a line that looks like a publication or known report name
    title = None
    for line in lines:
        if SKIP_TITLE_PATTERN.match(line):
            continue
        if PUBLICATION_NAME_PATTERN.match(line):
            title = line[:240]
            break
        if KNOWN_REPORT_PATTERN.match(line):
            title = line[:240]
            break

    # Tier 1.5: Try title-like lines only when publisher is unknown
    # (for known publishers, we fall through to Tier 2 for cleaner results)
    if title is None and publisher == "Unknown":
        body_kw = re.compile(
            r"(?i)(?:said|rose|fell|climbed|dropped|surged|plunged|reported|according|"
            r"percent|barrel|million|billion|traded|settled|closed|opened|index|stock|"
            r"share|bond|yield|treasury|federal|reserve|bank|economy|growth|data)"
        )
        for line in lines[:15]:
            if SKIP_TITLE_PATTERN.match(line):
                continue
            if TITLE_LIKE_PATTERN.match(line):
                if not body_kw.search(line):
                    if not re.search(r"(?:page|continued|volume|issue)", line, re.IGNORECASE):
                        title = line[:240]
                        break

    # Tier 2: Fallback to report_family if meaningful
    if title is None:
        rf = _detect_report_family(attachment.attachment_name, publisher)
        if rf and rf.lower() not in ("unknown", "platts", "general"):
            title = rf[:240]

    # Tier 3: Fallback — publisher name, or cleaned filename
    if title is None:
        if publisher != "Unknown":
            rf = _detect_report_family(attachment.attachment_name, publisher)
            if rf.lower() not in ("unknown", "platts", "general"):
                title = f"{publisher} {rf}"[:240]
            else:
                title = publisher[:240]
        else:
            name = attachment.attachment_name or path.stem
            name = re.sub(r"\.[^.]+$", "", name)
            name = re.sub(r"[_\-]\d{4,8}", "", name)
            name = re.sub(r"[_\-]+", " ", name).strip()
            title = name[:240] if len(name) >= 3 else path.stem[:240]

    # Always classify regions and commodities (was inside Tier 3 if-block)
    regions, commodities = _classify(parsed_text)
    sections = _split_sections(source_id, pages)
    for index, table in enumerate(tables):
        table.table_id = _stable_id("TBL", f"{source_id}:{index}")
        table.source_id = source_id
        matching_section = next((section for section in sections if section.page_start == table.page_number), None)
        if matching_section:
            table.section_id = matching_section.section_id
            matching_section.table_ids.append(table.table_id)

    raw_text_path: str | None = None
    if parsed_text and parsed_text_dir:
        parsed_text_dir.mkdir(parents=True, exist_ok=True)
        output_path = parsed_text_dir / f"{source_id}.txt"
        output_path.write_text(parsed_text, encoding="utf-8")
        raw_text_path = str(output_path)

    status = DocumentProcessingStatus.FAILED if error_message else (
        DocumentProcessingStatus.NEEDS_REVIEW if review_reasons else DocumentProcessingStatus.PARSED
    )
    return SourceDocument(
        parser_version=parser_version, source_id=source_id,
        ingestion=SourceIngestion(
            source_channel=telegram_input.source_channel, chat_id=telegram_input.message.telegram_chat_id,
            message_id=telegram_input.message.telegram_message_id,
            message_date=telegram_input.message.telegram_message_date, attachment_id=attachment_id,
        ),
        document=SourceDocumentMetadata(
            publisher=publisher, publisher_confidence=publisher_confidence,
            report_family=report_family, report_title=title,
            document_type="market_report", published_at=published_at, market_date=selected_date.value,
            market_date_confidence=selected_date.confidence,
            market_date_reason=market_date_reason,
            date_candidates=date_candidates, language="en", regions=regions,
            commodities=commodities, content_hash=attachment.attachment_hash,
        ),
        content=SourceContent(
            raw_text_path=raw_text_path, parsed_text=parsed_text, sections=sections, tables=tables,
            parse_method=method, parse_confidence=0.2 if error_message else (0.45 if review_reasons else 0.95),
        ),
        status=SourceStatus(
            processing_status=status, source_verified=publisher_confidence >= 0.8,
            needs_review=status != DocumentProcessingStatus.PARSED,
            review_reasons=review_reasons, error_message=error_message,
        ),
    )
